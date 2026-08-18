from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from sklearn.metrics import confusion_matrix

from src.evaluate import metrics as existing_metrics


ROOT = Path(__file__).resolve().parent
SIM_ROOT = ROOT / "illustrative_simulation"
PREDICTIONS = SIM_ROOT / "results" / "oof_predictions.csv"
SUMMARY = SIM_ROOT / "results" / "summary_metrics.csv"
AUDIT_CSV = SIM_ROOT / "results" / "confusion_matrix_and_metric_audit.csv"
OUTPUT = SIM_ROOT / "report" / "实验一_各组混淆矩阵与分类指标核对.docx"

ORDER = [
    "fusion_full", "ppg_only", "fusion_drop_demographic_lifestyle",
    "fusion_drop_procedure_medication", "fusion_drop_laboratory_echo_ecg",
    "fusion_drop_af_history", "clinical_only",
]
LABELS = {
    "fusion_full": "个人先验+纵向PPG",
    "ppg_only": "仅纵向PPG",
    "fusion_drop_demographic_lifestyle": "删除人口学及生活方式",
    "fusion_drop_procedure_medication": "删除手术及用药",
    "fusion_drop_laboratory_echo_ecg": "删除实验室、超声和心电",
    "fusion_drop_af_history": "删除房颤及既往病史",
    "clinical_only": "仅个人先验",
}


def divide(numerator, denominator):
    return numerator / denominator if denominator else 0.0


def compute_audit(predictions):
    rows = []
    for condition in ORDER:
        frame = predictions[predictions.condition == condition].sort_values("patient_id")
        y = frame.label.to_numpy(dtype=int)
        p = frame.probability.to_numpy(dtype=float)
        thresholds = frame.threshold.to_numpy(dtype=float)
        predicted = (p >= thresholds).astype(int)
        tn, fp, fn, tp = confusion_matrix(y, predicted, labels=[0, 1]).ravel()
        manual = {
            "sensitivity": divide(tp, tp + fn),
            "specificity": divide(tn, tn + fp),
            "accuracy": divide(tp + tn, len(y)),
            "precision": divide(tp, tp + fp),
            "npv": divide(tn, tn + fn),
            "f1": divide(2 * tp, 2 * tp + fp + fn),
            "brier": float(np.mean((p - y) ** 2)),
        }
        reference = existing_metrics(y, p, thresholds)
        checked = ["sensitivity", "specificity", "accuracy", "f1", "brier"]
        max_difference = max(abs(manual[name] - reference[name]) for name in checked)
        rows.append({"condition": condition, "condition_cn": LABELS[condition], "n": len(y),
                     "positive_n": int((y == 1).sum()), "negative_n": int((y == 0).sum()),
                     "tp": int(tp), "fn": int(fn), "tn": int(tn), "fp": int(fp),
                     **manual, "max_difference_vs_existing_code": max_difference,
                     "verification": "通过" if max_difference < 1e-12 else "不通过"})
    return pd.DataFrame(rows)


def set_cell_fill(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_table_geometry(table, widths_inches):
    if abs(sum(widths_inches) - 6.5) > 1e-9:
        raise ValueError("table widths must sum to 6.5 inches")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell, width in zip(row.cells, widths_inches):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(round(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    for grid_col, width in zip(table._tbl.tblGrid.gridCol_lst, widths_inches):
        grid_col.set(qn("w:w"), str(round(width * 1440)))


def style_table(table, widths):
    set_table_geometry(table, widths)
    set_repeat_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            if row_index == 0:
                set_cell_fill(cell, "F2F4F7")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if column_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(8)
                    run.font.bold = row_index == 0


def add_table(doc, headers, records, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for record in records:
        cells = table.add_row().cells
        for cell, value in zip(cells, record):
            cell.text = str(value)
    style_table(table, widths)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, before, after in (("Heading 1", 16, 16, 8), ("Heading 2", 13, 12, 6)):
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(46, 116, 181)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def build_report():
    predictions = pd.read_csv(PREDICTIONS)
    audit = compute_audit(predictions)
    audit.to_csv(AUDIT_CSV, index=False, encoding="utf-8-sig")

    summary = pd.read_csv(SUMMARY)
    point = summary.pivot(index="condition", columns="metric", values="estimate")
    for _, row in audit.iterrows():
        condition = row.condition
        for metric in ["sensitivity", "specificity", "accuracy", "f1", "brier"]:
            if abs(row[metric] - point.loc[condition, metric]) >= 1e-12:
                raise AssertionError(f"summary mismatch: {condition}/{metric}")

    doc = Document()
    configure_document(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("实验一：各组混淆矩阵与分类指标核对")
    run.font.name = "Microsoft YaHei"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 77, 120)
    subtitle = doc.add_paragraph("示意性模拟结果 | 基于逐患者折外预测重新计算")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. 核对方法", level=1)
    doc.add_paragraph(
        "对每个实验条件，按该患者对应的threshold字段将预测概率转换为二分类结果。随后直接统计真阳性（TP）、"
        "假阴性（FN）、真阴性（TN）和假阳性（FP），并由混淆矩阵重新计算各分类指标。"
        "同一结果再与src/evaluate.py中的现有评价函数以及summary_metrics.csv中的点估计逐项比较。"
    )
    doc.add_paragraph(
        "敏感度=TP/(TP+FN)；特异度=TN/(TN+FP)；准确率=(TP+TN)/N；精确率=TP/(TP+FP)；"
        "阴性预测值=TN/(TN+FN)；F1=2TP/(2TP+FP+FN)；Brier=N^{-1}Σ(p-y)^2。"
    )

    doc.add_heading("2. 混淆矩阵及敏感度、特异度", level=1)
    records = []
    for _, row in audit.iterrows():
        records.append([row.condition_cn, row.n, row.tp, row.fn, row.tn, row.fp,
                        f"{row.sensitivity:.3f}", f"{row.specificity:.3f}"])
    add_table(doc, ["实验条件", "N", "TP", "FN", "TN", "FP", "敏感度", "特异度"],
              records, [2.35, 0.45, 0.45, 0.45, 0.45, 0.45, 0.95, 0.95])
    p = doc.add_paragraph("表1  各实验条件的混淆矩阵与核心分类指标")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("3. 补充分类指标及一致性检查", level=1)
    records = []
    for _, row in audit.iterrows():
        records.append([row.condition_cn, f"{row.accuracy:.3f}", f"{row.precision:.3f}",
                        f"{row.npv:.3f}", f"{row.f1:.3f}", f"{row.brier:.3f}",
                        f"{row.max_difference_vs_existing_code:.2e}", row.verification])
    add_table(doc, ["实验条件", "准确率", "精确率", "NPV", "F1", "Brier", "最大差值", "核对"],
              records, [1.9, 0.72, 0.72, 0.65, 0.6, 0.65, 0.78, 0.48])
    p = doc.add_paragraph("表2  补充指标及与现有评价代码的一致性")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    max_delta = audit.max_difference_vs_existing_code.max()
    doc.add_heading("4. 核对结论", level=1)
    doc.add_paragraph(
        f"7个实验条件均包含{int(audit.n.iloc[0])}例患者，其中阳性{int(audit.positive_n.iloc[0])}例、"
        f"阴性{int(audit.negative_n.iloc[0])}例。独立公式与现有评价函数之间的最大绝对差值为{max_delta:.2e}。"
        "敏感度、特异度、准确率、F1和Brier与summary_metrics.csv中的点估计完全一致，未发现计算错误。"
    )
    doc.add_paragraph("本文件核对的是示意性模拟结果，不代表真实临床队列中的模型效能。")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT, audit


if __name__ == "__main__":
    output, audit = build_report()
    print(output)
    print(audit[["condition", "tp", "fn", "tn", "fp", "sensitivity", "specificity", "verification"]].to_string(index=False))
