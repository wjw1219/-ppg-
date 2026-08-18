from pathlib import Path
import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
RESULTS, FIGURES, REPORT = ROOT / "results", ROOT / "figures", ROOT / "report"
REPORT.mkdir(exist_ok=True)
OUTPUT = REPORT / "实验一_信息来源消融_实验说明.docx"
summary = pd.read_csv(RESULTS / "summary_metrics.csv")
paired = pd.read_csv(RESULTS / "paired_differences.csv")

LABELS = {
    "clinical_only": "仅个人先验", "ppg_only": "仅纵向PPG", "fusion_full": "个人先验+纵向PPG",
    "fusion_drop_demographic_lifestyle": "删除人口学及生活方式",
    "fusion_drop_af_history": "删除房颤及既往病史",
    "fusion_drop_laboratory_echo_ecg": "删除实验室、超声和心电",
    "fusion_drop_procedure_medication": "删除手术及用药",
}

def shade(cell, fill):
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); cell._tc.get_or_add_tcPr().append(shd)

def cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr(); tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None: tc_w = OxmlElement("w:tcW"); tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width)); tc_w.set(qn("w:type"), "dxa")

def table_style(table, widths, size=8.5):
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr; tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None: tbl_w = OxmlElement("w:tblW"); tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths))); tbl_w.set(qn("w:type"), "dxa")
    ind = OxmlElement("w:tblInd"); ind.set(qn("w:w"), "120"); ind.set(qn("w:type"), "dxa"); tbl_pr.append(ind)
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for i, row in enumerate(table.rows):
        for cell, width in zip(row.cells, widths):
            cell_width(cell, width); cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if i == 0: shade(cell, "F2F4F7")
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = "Calibri"; run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(size); run.bold = i == 0

def mrow(condition, name):
    return summary[(summary.condition == condition) & (summary.metric == name)].iloc[0]

def ci(condition, name):
    r = mrow(condition, name); return f"{r.estimate:.3f} ({r.ci_lower:.3f}-{r.ci_upper:.3f})"

def performance_table(doc, conditions):
    headers = ["模型/条件", "ROC-AUC (95% CI)", "AUPRC (95% CI)", "敏感度", "特异度", "F1", "Brier"]
    table = doc.add_table(rows=1, cols=7); table.style = "Table Grid"
    for c, h in zip(table.rows[0].cells, headers): c.text = h
    for condition in conditions:
        vals = [LABELS[condition], ci(condition,"roc_auc"), ci(condition,"auprc"),
                f"{mrow(condition,'sensitivity').estimate:.3f}", f"{mrow(condition,'specificity').estimate:.3f}",
                f"{mrow(condition,'f1').estimate:.3f}", f"{mrow(condition,'brier').estimate:.3f}"]
        for c, v in zip(table.add_row().cells, vals): c.text = v
    table_style(table, [1900,1600,1600,1050,1050,980,1180], 8)

doc = Document(); section = doc.sections[0]
section.page_width = Inches(8.5); section.page_height = Inches(11)
section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(.492)
normal = doc.styles["Normal"]; normal.font.name = "Calibri"; normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei"); normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in [("Heading 1",16,"2E74B5",16,8),("Heading 2",13,"2E74B5",12,6),("Heading 3",12,"1F4D78",8,4)]:
    s = doc.styles[name]; s.font.name = "Calibri"; s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    s.font.size = Pt(size); s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before = Pt(before); s.paragraph_format.space_after = Pt(after); s.paragraph_format.keep_with_next = True
header = section.header.paragraphs[0]; header.text = "实验一 | 信息来源消融"
for r in header.runs: r.font.size = Pt(9); r.font.color.rgb = RGBColor(100,100,100)
footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT; footer.add_run("第 ")
field = OxmlElement("w:fldSimple"); field.set(qn("w:instr"), "PAGE"); footer._p.append(field); footer.add_run(" 页")

p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(4)
r = p.add_run("实验一：信息来源消融"); r.bold = True; r.font.size = Pt(23); r.font.color.rgb = RGBColor(31,78,121)
r.font.name = "Calibri"; r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(14)
r = p.add_run("7天PPG主尺度 | 完整26周观察 | 患者级五折交叉验证"); r.font.size = Pt(12); r.font.color.rgb = RGBColor(90,90,90)
for label, value in [("分析对象","387例患者"),("比较条件","3种主要输入条件及4项个人先验消融"),("主要指标","ROC-AUC"),("运行日期","2026年7月23日")]:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2); a = p.add_run(label + "："); a.bold = True; p.add_run(value)

doc.add_heading("1. 实验目的", 1)
doc.add_paragraph("本实验在PPG窗口固定为预设7天主分析尺度、使用完整26周纵向数据的条件下，比较仅个人先验、仅纵向PPG以及个人先验与纵向PPG表示层融合的预测性能。随后，在完整融合模型中依次删除四类个人先验信息，以评价各类先验信息对模型性能的边际影响。")
doc.add_heading("2. 数据与结局", 1)
doc.add_paragraph("本实验使用387例模拟患者数据验证实验流程。每位患者包含出院时个人先验信息、26个7天PPG时间窗口以及出院和6个月随访时的72小时Holter房颤负荷。Holter房颤负荷仅用于确定6个月负荷相较出院时是否增加，不进入任何模型。")
table = doc.add_table(rows=1, cols=2); table.style = "Table Grid"; table.rows[0].cells[0].text = "项目"; table.rows[0].cells[1].text = "结果"
for a,b in [("总患者数","387"),("房颤负荷增加","200 (51.7%)"),("房颤负荷未增加","187 (48.3%)"),("7天PPG窗口","10,062"),("每位患者时间节点","26")]:
    cells = table.add_row().cells; cells[0].text = a; cells[1].text = b
table_style(table, [2800,6560], 9)

doc.add_heading("3. 输入信息与消融条件", 1)
groups = [("人口学及生活方式","年龄、性别、BMI、吸烟、饮酒"),("房颤及既往病史","房颤类型、房颤病程、高血压、心力衰竭、糖尿病、卒中或TIA史、既往消融次数"),("实验室、超声和心电","hs-CRP、NT-proBNP、LAD、LAVI、LVEF；当前数据没有独立基线心电字段"),("手术及用药","消融能量、出院抗心律失常药、出院抗凝药")]
table = doc.add_table(rows=1, cols=2); table.style = "Table Grid"; table.rows[0].cells[0].text = "信息类别"; table.rows[0].cells[1].text = "变量"
for a,b in groups: cells = table.add_row().cells; cells[0].text = a; cells[1].text = b
table_style(table, [2600,6760], 9)

doc.add_heading("4. 模型与融合方法", 1)
doc.add_paragraph("临床分支将变量按上述四类关系分别编码，通过关系特异的变量注意力汇总得到患者级临床表示。PPG分支将每周PPG特征构建为26个时间节点，仅设置相邻周双向边和自环，通过时间图注意力更新节点，并采用掩码注意力池化得到患者级PPG表示。")
doc.add_paragraph("完整融合模型仅采用表示层融合：将临床表示与PPG表示直接拼接，再输入两层多层感知机输出房颤负荷增加概率。不使用门控融合、概率平均或节点级临床-PPG连接。")
doc.add_heading("5. 训练与统计方法", 1)
doc.add_paragraph("所有条件共享相同的患者级分层五折划分。每个外层训练折内部按标签分层划分20%验证集，用于早停和Youden指数阈值确定。连续变量填补和标准化、分类变量编码以及PPG标准化均仅在相应训练数据中拟合。每个外层折使用3个固定随机种子训练，测试患者的3个概率取平均。")
doc.add_paragraph("主要指标为ROC-AUC，次要指标包括AUPRC、敏感度、特异度、F1、准确率和Brier score。95%置信区间采用2,000次分层患者级Bootstrap计算；完整融合模型与其他条件的差值采用相同重采样患者索引进行配对计算。")

doc.add_heading("6. 主要模态比较", 1)
performance_table(doc, ["clinical_only","ppg_only","fusion_full"])
doc.add_paragraph("仅纵向PPG取得最高ROC-AUC（0.593，95% CI 0.534-0.649）和AUPRC（0.596，95% CI 0.545-0.656）。完整融合模型的ROC-AUC为0.572，未超过仅PPG模型；仅个人先验模型的ROC-AUC为0.565。因此，本次实验没有观察到表示层融合带来的增益。")
doc.add_picture(str(FIGURES / "roc_curves.png"), width=Inches(6.1)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("图1. 三种主要信息来源模型的折外ROC曲线"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].italic = True

doc.add_heading("7. 个人先验内部消融", 1)
performance_table(doc, ["fusion_full","fusion_drop_demographic_lifestyle","fusion_drop_af_history","fusion_drop_laboratory_echo_ecg","fusion_drop_procedure_medication"])
doc.add_paragraph("四项删除条件的ROC-AUC均低于完整融合模型，其中删除房颤及既往病史后的ROC-AUC最低（0.513）。但是，各模型置信区间重叠较多，因此这些结果只能说明删除相应信息后预测性能在本次数据中出现下降，不能据此推断变量类别具有因果作用。")
doc.add_picture(str(FIGURES / "ablation_performance.png"), width=Inches(6.2)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("图2. 完整融合模型与四项个人先验消融的ROC-AUC"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].italic = True

doc.add_heading("8. 配对性能差值", 1)
table = doc.add_table(rows=1, cols=3); table.style = "Table Grid"
for c,h in zip(table.rows[0].cells,["完整融合减比较条件","ROC-AUC差值 (95% CI)","AUPRC差值 (95% CI)"]): c.text = h
for condition in LABELS:
    if condition == "fusion_full": continue
    rows = paired[paired.comparison == f"fusion_full - {condition}"].set_index("metric")
    cells = table.add_row().cells; cells[0].text = LABELS[condition]
    for cell, name in zip(cells[1:], ["roc_auc_difference","auprc_difference"]):
        row = rows.loc[name]; cell.text = f"{row.estimate:+.3f} ({row.ci_lower:+.3f}至{row.ci_upper:+.3f})"
table_style(table, [3600,2880,2880], 8.5)

doc.add_heading("9. 结果解释", 1)
doc.add_paragraph("本次实验中，纵向PPG包含与房颤负荷变化相关的时间动态信息，因此仅PPG模型表现最好。临床先验分支的判别能力有限，将其表示直接拼接到PPG表示后增加了模型参数和估计方差，但没有提供足够的互补信息，可能是融合模型未获得增益的原因。")
doc.add_paragraph("先验消融结果表明，删除四类信息后性能点估计均下降，其中房颤及既往病史、手术及用药删除后的下降相对明显。不过，差值是否稳定应以配对Bootstrap置信区间为准，不应只依据点估计排序。")
doc.add_heading("10. 实验限制", 1)
for text in ["样本量为387例，深度模型估计方差较大。","本实验只评价7天尺度和完整26周输入，不能回答不同观察长度或聚合尺度的问题。","实验室、超声和心电类别中没有独立基线心电字段，因此该消融实际删除实验室与超声变量。","超参数在所有条件间固定以保证公平比较，但未进行嵌套交叉验证下的大范围搜索。"]:
    doc.add_paragraph(text, style="List Bullet")
doc.add_heading("11. 可复现输出", 1)
doc.add_paragraph("实验目录保存逐患者折外预测、逐折训练记录、总体指标、配对差值、质量检查、ROC图和消融图。所有结果均可从results/oof_predictions.csv重新计算。")
doc.save(OUTPUT); print(OUTPUT)
