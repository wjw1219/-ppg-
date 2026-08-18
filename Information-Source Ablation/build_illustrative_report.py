from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
TARGET = ROOT / "illustrative_simulation"
RESULTS, FIGURES, REPORT = TARGET / "results", TARGET / "figures", TARGET / "report"
OUTPUT = REPORT / "实验一_信息来源消融_示意性模拟结果说明.docx"

LABELS = {
    "fusion_full": "个人先验+纵向PPG",
    "ppg_only": "仅纵向PPG",
    "fusion_drop_demographic_lifestyle": "删除人口学及生活方式",
    "fusion_drop_procedure_medication": "删除手术及用药",
    "fusion_drop_laboratory_echo_ecg": "删除实验室、超声和心电",
    "fusion_drop_af_history": "删除房颤及既往病史",
    "clinical_only": "仅个人先验",
}
ORDER = list(LABELS)
METRICS = ["roc_auc", "auprc", "sensitivity", "specificity", "f1", "brier"]
HEADERS = ["实验条件", "ROC-AUC (95% CI)", "AUPRC (95% CI)", "敏感度", "特异度", "F1", "Brier"]


def format_ci(estimate, lower, upper):
    return f"{estimate:.3f} ({lower:.3f}-{upper:.3f})"


def set_cell_fill(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = Inches(0.8)
    section.left_margin = section.right_margin = Inches(0.75)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size in (("Title", 20), ("Heading 1", 15), ("Heading 2", 12)):
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(46, 116, 181)
        style.font.bold = True


def add_metric_table(doc, summary):
    indexed = summary.set_index(["condition", "metric"])
    table = doc.add_table(rows=1, cols=len(HEADERS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(1.55), Inches(1.25), Inches(1.25), Inches(0.72), Inches(0.72), Inches(0.62), Inches(0.72)]
    for cell, header, width in zip(table.rows[0].cells, HEADERS, widths):
        cell.width = width
        cell.text = header
        set_cell_fill(cell, "E8EEF5")
    for condition in ORDER:
        cells = table.add_row().cells
        cells[0].text = LABELS[condition]
        for column, metric in enumerate(METRICS, start=1):
            row = indexed.loc[(condition, metric)]
            cells[column].text = format_ci(row.estimate, row.ci_lower, row.ci_upper)
        for cell, width in zip(cells, widths):
            cell.width = width
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if column_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run.font.size = Pt(7.5 if row_index else 7.2)
                    run.font.bold = row_index == 0
    return table


def build_report():
    summary = pd.read_csv(RESULTS / "summary_metrics.csv")
    REPORT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("实验一：信息来源消融\n示意性模拟结果说明")
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("本文件仅用于展示预期结果形式，不是实际模型效能结果。")
    run.bold = True
    run.font.color.rgb = RGBColor(155, 28, 28)

    doc.add_heading("1. 模拟目的与计算方法", level=1)
    doc.add_paragraph(
        "本次示意性模拟用于展示信息来源消融实验的结果组织方式。程序首先在患者层面生成各实验条件的预测概率，"
        "再使用与实验一相同的评价函数计算指标。因此，ROC-AUC、AUPRC、敏感度、特异度、F1和Brier均来源于同一组逐患者预测，未直接填写汇总数值。"
    )
    doc.add_paragraph(
        "ROC-AUC和AUPRC直接根据真实标签与预测概率计算；敏感度、特异度和F1根据0.5分类阈值得到；"
        "Brier分数为预测概率与二分类标签之间的均方误差。95%置信区间采用2000次分层Bootstrap的2.5和97.5百分位数。"
    )

    doc.add_heading("2. 示意性性能结果", level=1)
    doc.add_paragraph("表1列出7种实验条件的示意性评价结果。括号内为分层Bootstrap 95%置信区间，所有指标保留三位小数。")
    caption = doc.add_paragraph("表1  信息来源消融的示意性模拟结果")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].bold = True
    add_metric_table(doc, summary)
    doc.add_paragraph(
        "注：Brier分数越低越好；其余指标越高越好。所有数值均为示意性模拟结果，不代表真实队列中的模型表现。"
    )

    doc.add_heading("3. 图形结果", level=1)
    doc.add_paragraph("图1比较仅个人先验、仅纵向PPG和完整融合模型的ROC曲线。")
    doc.add_picture(str(FIGURES / "roc_curves.png"), width=Inches(5.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("图1  主要信息来源模型的示意性ROC曲线")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("图2显示完整融合模型及依次删除四类个人先验信息后的ROC-AUC。")
    doc.add_picture(str(FIGURES / "ablation_performance.png"), width=Inches(6.1))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("图2  个人先验内部消融的示意性性能比较")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("4. 结果解释", level=1)
    doc.add_paragraph(
        "示意结果预设完整融合模型性能最高，仅纵向PPG次之，仅个人先验最低。内部消融中，删除房颤及既往病史造成的性能下降最大，"
        "其次为实验室、超声和心电信息、手术及用药信息以及人口学及生活方式信息。该排序用于演示结果表达，不构成临床效应证据。"
    )
    doc.add_paragraph(
        "正式实验应完全依据真实折外预测重新运行make_results.py，并报告真实点估计和置信区间，不应将本文件中的示意性数值用于效能结论。"
    )
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_report())
