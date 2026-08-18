import json
import unittest
from pathlib import Path

import pandas as pd
from docx import Document

ROOT = Path(__file__).resolve().parents[1]


class OutputTests(unittest.TestCase):
    def test_oof_predictions_and_quality_checks(self):
        pred = pd.read_csv(ROOT / "results/oof_predictions.csv")
        self.assertEqual(len(pred), 2709)
        self.assertEqual(pred.condition.nunique(), 7)
        self.assertTrue(pred.groupby(["patient_id", "condition"]).size().eq(1).all())
        self.assertTrue(pred.probability.between(0, 1).all())
        checks = json.loads((ROOT / "results/quality_checks.json").read_text())
        self.assertTrue(all(checks.values()))

    def test_report_structure(self):
        path = ROOT / "report/实验一_信息来源消融_实验说明.docx"
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        self.assertEqual(len(doc.inline_shapes), 2)
        self.assertEqual(len(doc.tables), 5)
        self.assertEqual(sum(p.style.name == "Heading 1" for p in doc.paragraphs), 11)
        self.assertNotIn("TODO", text)
        self.assertNotIn("TBD", text)
        section = doc.sections[0]
        self.assertAlmostEqual(section.left_margin.inches, 1.0, places=2)
        self.assertAlmostEqual(section.right_margin.inches, 1.0, places=2)


if __name__ == "__main__": unittest.main()
