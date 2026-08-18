import json
import unittest
from pathlib import Path

import pandas as pd
from docx import Document


ROOT = Path(__file__).resolve().parents[1]


class OutputTests(unittest.TestCase):
    def test_result_cardinality_and_checks(self):
        self.assertEqual(len(pd.read_csv(ROOT / "results/weekly_oof_predictions.csv")), 10062)
        self.assertEqual(len(pd.read_csv(ROOT / "results/model_oof_predictions.csv")), 2322)
        self.assertEqual(len(pd.read_csv(ROOT / "results/fold_metrics.csv")), 465)
        self.assertTrue(all(json.loads((ROOT / "results/quality_checks.json").read_text()).values()))

    def test_report_structure(self):
        path = ROOT / "report/实验三_逐周减量与模型比较_实验说明.docx"
        self.assertTrue(path.exists())
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        for required in ["实验目的", "逐周减量结果", "完整26周模型比较", "计算成本", "模拟数据"]:
            self.assertIn(required, text)
        self.assertNotIn("TODO", text)
        self.assertNotIn("待补充", text)
        self.assertGreaterEqual(len(doc.tables), 5)
        self.assertEqual(len(doc.inline_shapes), 2)
        section = doc.sections[0]
        self.assertAlmostEqual(section.left_margin.inches, 1.0, places=2)
        self.assertAlmostEqual(section.right_margin.inches, 1.0, places=2)


if __name__ == "__main__":
    unittest.main()
