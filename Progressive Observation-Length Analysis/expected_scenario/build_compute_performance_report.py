import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(ROOT.parent))
from src.data import FoldPreprocessor, load_common, load_ppg
from src.models import ComparisonModel

EXP = ROOT.parent; RES = ROOT / "results"
INPUT = EXP.parent / "outputs/evoaf_synthetic_387/intermediate"
OUT = ROOT / "report" / "实验三_逐周与模型计算性能大表.docx"
NAMES={"mlp":"特征拼接+MLP","gru":"GRU","lstm":"LSTM","transformer":"Transformer","temporal_gat":"普通时序GAT","full_model":"完整模型"}

baseline, labels, _ = load_common(INPUT); ppg = load_ppg(INPUT); ids = baseline.index.tolist()
prep = FoldPreprocessor(26).fit(baseline, ppg, ids); sample = prep.transform(baseline, ppg, labels, ids[:1])
clinical_dim=sample.clinical.shape[1]; group_dims=[b-a for a,b in sample.group_slices]
resource=pd.read_csv(EXP/"results/resource_usage.csv").set_index("model")
weekly_metric=pd.read_csv(RES/"expected_weekly_metrics.csv"); model_metric=pd.read_csv(RES/"expected_model_metrics.csv")

def parameter_count(kind,weeks):
    model=ComparisonModel(kind,clinical_dim,group_dims,weeks,h=32,heads=2,drop=.25)
    return sum(x.numel() for x in model.parameters())

full_time=float(resource.loc["full_model","cpu_single_patient_ms_median"])
full_memory=float(resource.loc["full_model","process_rss_mb"])
weekly=[]
for week in range(1,27):
    ratio=week/26
    inference=0.72+(full_time-0.72)*(0.32*ratio+0.68*ratio**2)
    memory=230+(full_memory-230)*(0.42*ratio+0.58*ratio**2)
    auc=weekly_metric[(weekly_metric.weeks==week)&(weekly_metric.metric=="roc_auc")].iloc[0]
    weekly.append({"weeks":week,"parameters":parameter_count("full_model",week),"inference_time_ms":inference,"peak_memory_mb":memory,"roc_auc":auc.estimate,"roc_auc_low":auc.ci_lower,"roc_auc_high":auc.ci_upper,"measurement":"参数量实际计算；时间与内存为锚定26周CPU基准的模拟估计"})

models=[]
for kind in NAMES:
    auc=model_metric[(model_metric.model==kind)&(model_metric.metric=="roc_auc")].iloc[0]
    parameters=parameter_count(kind,26)
    if parameters != int(resource.loc[kind,"parameters"]): raise ValueError(f"parameter mismatch: {kind}")
    models.append({"model":kind,"parameters":parameters,"inference_time_ms":float(resource.loc[kind,"cpu_single_patient_ms_median"]),"peak_memory_mb":float(resource.loc[kind,"process_rss_mb"]),"roc_auc":auc.estimate,"roc_auc_low":auc.ci_lower,"roc_auc_high":auc.ci_upper,"measurement":"CPU基准测量"})

weekly_df=pd.DataFrame(weekly); model_df=pd.DataFrame(models)
if not weekly_df.inference_time_ms.is_monotonic_increasing or not weekly_df.peak_memory_mb.is_monotonic_increasing: raise ValueError("weekly compute cost must increase")
if weekly_df.iloc[-1].parameters != int(resource.loc["full_model","parameters"]): raise ValueError("week-26 parameters mismatch")
weekly_df.to_csv(RES/"weekly_compute_performance.csv",index=False); model_df.to_csv(RES/"model_compute_performance.csv",index=False)

doc=Document(); sec=doc.sections[0]; sec.orientation=WD_ORIENT.LANDSCAPE; sec.page_width=Inches(11); sec.page_height=Inches(8.5); sec.top_margin=sec.bottom_margin=Inches(.65); sec.left_margin=sec.right_margin=Inches(.7)
normal=doc.styles["Normal"]; normal.font.name="Microsoft YaHei"; normal.font.size=Pt(9.5); normal.paragraph_format.space_after=Pt(5)
for style,size in [("Heading 1",15),("Heading 2",12)]: doc.styles[style].font.name="Microsoft YaHei"; doc.styles[style].font.size=Pt(size); doc.styles[style].font.color.rgb=RGBColor(31,78,121)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("实验三：逐周与模型计算性能大表"); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=RGBColor(31,78,121)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("预期结果模拟 | CPU环境 | GPU显存未测量"); r.bold=True; r.font.color.rgb=RGBColor(155,28,28)
doc.add_heading("1. 指标口径",1)
doc.add_paragraph("Parameters为当前PyTorch模型结构的可训练参数总数。六模型Inference time为预热后、单患者、批量大小1的CPU推理时间中位数；Peak memory为包含Python与PyTorch运行时的进程RSS。逐周Inference time和Peak memory不是重新实测值，而是以第26周完整模型CPU基准为锚点，结合序列长度及时间注意力计算量生成的单调递增模拟估计。")

def add_table(title,headers,rows):
    doc.add_heading(title,1); table=doc.add_table(rows=1,cols=len(headers)); table.style="Table Grid"; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    for c,x in zip(table.rows[0].cells,headers): c.text=x
    for values in rows:
        for c,x in zip(table.add_row().cells,values): c.text=str(x)
    for i,row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
                for run in p.runs: run.font.name="Microsoft YaHei"; run.font.size=Pt(8); run.bold=i==0
    return table

add_table("2. 完整模型前1至前26周计算性能",["周数","Parameters","Inference time/ms","Peak memory/MB","ROC-AUC (95% CI)"] ,[(r.weeks,f"{int(r.parameters):,}",f"{r.inference_time_ms:.2f}",f"{r.peak_memory_mb:.1f}",f"{r.roc_auc:.3f} ({r.roc_auc_low:.3f}-{r.roc_auc_high:.3f})") for r in weekly_df.itertuples()])
add_table("3. 完整26周六模型计算性能",["模型","Parameters","Inference time/ms","Peak memory/MB","ROC-AUC (95% CI)"] ,[(NAMES[r.model],f"{int(r.parameters):,}",f"{r.inference_time_ms:.2f}",f"{r.peak_memory_mb:.1f}",f"{r.roc_auc:.3f} ({r.roc_auc_low:.3f}-{r.roc_auc_high:.3f})") for r in model_df.itertuples()])
doc.add_heading("4. 综合解释",1)
doc.add_paragraph("MLP推理最快，但预测性能较低；Transformer参数量、推理时间和内存开销均较高；GRU、LSTM和普通时序GAT处于中间水平。完整模型的参数量低于MLP、LSTM和Transformer，CPU推理时间低于Transformer且与LSTM接近，同时获得最高ROC-AUC，因此在本预期模拟中表现为预测性能与计算成本之间的综合最佳方案。该结论仅用于实验展示，真实计算成本需在最终部署硬件上重复测量。")
doc.save(OUT); print(OUT)
