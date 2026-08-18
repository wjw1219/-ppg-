import sys
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(ROOT.parent))
from src.evaluate import metrics

RES = ROOT / "results"
REPORT = ROOT / "report" / "实验三_混淆矩阵与指标公式核对.docx"
WEEKLY_PRED = pd.read_csv(RES / "expected_weekly_predictions.csv")
MODEL_PRED = pd.read_csv(RES / "expected_model_predictions.csv")
WEEKLY_METRIC = pd.read_csv(RES / "expected_weekly_metrics.csv")
MODEL_METRIC = pd.read_csv(RES / "expected_model_metrics.csv")
MODEL_NAMES = {"mlp":"特征拼接+MLP","gru":"GRU","lstm":"LSTM","transformer":"Transformer","temporal_gat":"普通时序GAT","full_model":"完整模型"}
METRIC_NAMES = ["roc_auc","auprc","sensitivity","specificity","f1","accuracy","brier"]

def audit_group(frame, key, value, stored):
    data = frame[frame[key] == value].sort_values("patient_id")
    y = data.label.to_numpy(int); p = data.probability.to_numpy(float)
    threshold = np.full(len(y), .5); pred = (p >= threshold).astype(int)
    tp = int(((pred == 1) & (y == 1)).sum()); fn = int(((pred == 0) & (y == 1)).sum())
    tn = int(((pred == 0) & (y == 0)).sum()); fp = int(((pred == 1) & (y == 0)).sum())
    recalculated = metrics(y, p, threshold)
    saved = stored[stored[key] == value].set_index("metric")
    differences = {name: abs(recalculated[name] - float(saved.loc[name,"estimate"])) for name in METRIC_NAMES}
    if max(differences.values()) > 1e-12:
        raise ValueError(f"{key}={value} metric mismatch: {differences}")
    precision = tp / max(tp + fp, 1)
    return {key:value,"n":len(y),"positive":int(y.sum()),"negative":int((y==0).sum()),
            "tp":tp,"fn":fn,"tn":tn,"fp":fp,"precision":precision,**recalculated,
            "max_difference":max(differences.values())}

weekly_rows = [audit_group(WEEKLY_PRED,"weeks",week,WEEKLY_METRIC) for week in range(1,27)]
model_rows = [audit_group(MODEL_PRED,"model",model,MODEL_METRIC) for model in MODEL_NAMES]
pd.DataFrame(weekly_rows).to_csv(RES / "weekly_confusion_matrix_audit.csv",index=False)
pd.DataFrame(model_rows).to_csv(RES / "model_confusion_matrix_audit.csv",index=False)

doc = Document(); section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE; section.page_width = Inches(11); section.page_height = Inches(8.5)
section.top_margin = section.bottom_margin = Inches(.65); section.left_margin = section.right_margin = Inches(.6)
normal=doc.styles["Normal"]; normal.font.name="Microsoft YaHei"; normal.font.size=Pt(9); normal.paragraph_format.space_after=Pt(4)
for style,size in [("Heading 1",15),("Heading 2",12)]: doc.styles[style].font.name="Microsoft YaHei"; doc.styles[style].font.size=Pt(size); doc.styles[style].font.color.rgb=RGBColor(31,78,121)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("实验三：混淆矩阵与指标公式核对"); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=RGBColor(31,78,121)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("预期结果模拟数据 | 固定分类阈值0.5 | 非实际模型训练结果"); r.bold=True; r.font.color.rgb=RGBColor(155,28,28)
doc.add_heading("1. 核对规则",1)
doc.add_paragraph("对每个周数和每个模型，使用患者级标签与预测概率重新构建混淆矩阵。预测概率大于或等于0.5记为房颤负荷增加，否则记为未增加。所有点估计均重新调用正式实验三的src/evaluate.py计算，并与指标CSV逐项比较；允许的最大绝对误差为1×10⁻¹²。")
doc.add_heading("2. 计算公式",1)
for text in ["敏感度 = TP / (TP + FN)","特异度 = TN / (TN + FP)","精确率 = TP / (TP + FP)","F1 = 2TP / (2TP + FP + FN)","准确率 = (TP + TN) / N","Brier = Σ(pᵢ - yᵢ)² / N"]: doc.add_paragraph(text,style="List Bullet")
doc.add_paragraph("ROC-AUC和AUPRC使用连续预测概率计算，不依赖0.5分类阈值。95%置信区间使用2,000次分层患者Bootstrap；每次重采样后均重新计算相应指标。")

def add_table(title, rows, key, label):
    doc.add_heading(title,1)
    headers=[label,"N","阳性","阴性","TP","FN","TN","FP","敏感度","特异度","精确率","F1","准确率","Brier","最大误差"]
    table=doc.add_table(rows=1,cols=len(headers)); table.style="Table Grid"; table.alignment=WD_TABLE_ALIGNMENT.CENTER
    for c,x in zip(table.rows[0].cells,headers): c.text=x
    for row in rows:
        values=[row[key],row["n"],row["positive"],row["negative"],row["tp"],row["fn"],row["tn"],row["fp"],f'{row["sensitivity"]:.3f}',f'{row["specificity"]:.3f}',f'{row["precision"]:.3f}',f'{row["f1"]:.3f}',f'{row["accuracy"]:.3f}',f'{row["brier"]:.3f}',f'{row["max_difference"]:.1e}']
        for c,x in zip(table.add_row().cells,values): c.text=str(MODEL_NAMES.get(x,x))
    for i,row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER; paragraph.paragraph_format.space_before=Pt(0); paragraph.paragraph_format.space_after=Pt(0)
                for run in paragraph.runs: run.font.name="Microsoft YaHei"; run.font.size=Pt(6.8); run.bold=i==0
    table.rows[0]._tr.get_or_add_trPr().append(__import__('docx').oxml.OxmlElement("w:tblHeader"))

add_table("3. 前1至前26周混淆矩阵与重算指标",weekly_rows,"weeks","周数")
add_table("4. 完整26周模型比较的混淆矩阵与重算指标",model_rows,"model","模型")
doc.add_heading("5. 核对结论",1)
doc.add_paragraph("全部周数和全部模型的重算点估计与当前指标CSV完全一致，最大绝对误差不超过1×10⁻¹²。敏感度、特异度、精确率、F1和准确率均由本报告所列TP、FN、TN和FP按公式计算；Brier、ROC-AUC和AUPRC均由患者级预测概率直接计算。")
doc.save(REPORT); print(REPORT)
