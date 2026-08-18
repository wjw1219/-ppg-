from pathlib import Path
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT=Path(__file__).resolve().parent; RES=ROOT/"results"; FIG=ROOT/"figures"; OUT=ROOT/"report/实验三_预期结果模拟_V5_指标核对版.docx"; OUT.parent.mkdir(exist_ok=True)
w=pd.read_csv(RES/"expected_weekly_metrics.csv"); m=pd.read_csv(RES/"expected_model_metrics.csv")
names={"mlp":"特征拼接+MLP","gru":"GRU","lstm":"LSTM","transformer":"Transformer","temporal_gat":"普通时序GAT","full_model":"完整模型"}
def val(df,key,item,metric): return df[(df[key]==item)&(df.metric==metric)].iloc[0]
def ci(df,key,item,metric):
    r=val(df,key,item,metric); return f"{r.estimate:.3f} ({r.ci_lower:.3f}-{r.ci_upper:.3f})"
doc=Document(); sec=doc.sections[0]; sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
n=doc.styles["Normal"]; n.font.name="Microsoft YaHei"; n.font.size=Pt(10.5); n.paragraph_format.space_after=Pt(6)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("实验三：预期结果模拟"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=RGBColor(155,28,28)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("非实际模型训练结果，不得作为临床效能证据"); r.bold=True; r.font.size=Pt(14); r.font.color.rgb=RGBColor(155,28,28)
doc.add_heading("1. 模拟目的",1); doc.add_paragraph("本文件用于展示预期结果的统计表和论文图形版式。患者标签和预测概率均由情景脚本生成，并非由真实数据或模型训练获得。情景预设完整模型性能最高，并预设逐周模型从第16周开始进入性能平台。")
doc.add_heading("2. 逐周预期性能",1); doc.add_picture(str(FIG/"expected_weekly_performance.png"),width=Inches(6.2)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
t=doc.add_table(rows=1,cols=7); t.style="Table Grid"; hdr=["周数","ROC-AUC (95% CI)","AUPRC (95% CI)","敏感度","特异度","F1","Brier"]
for c,x in zip(t.rows[0].cells,hdr): c.text=x
for week in range(1,27):
    row=[week,ci(w,"weeks",week,"roc_auc"),ci(w,"weeks",week,"auprc"),f"{val(w,'weeks',week,'sensitivity').estimate:.3f}",f"{val(w,'weeks',week,'specificity').estimate:.3f}",f"{val(w,'weeks',week,'f1').estimate:.3f}",f"{val(w,'weeks',week,'brier').estimate:.3f}"]
    for c,x in zip(t.add_row().cells,row): c.text=str(x)
doc.add_heading("3. 模型对比预期性能",1); doc.add_picture(str(FIG/"expected_model_comparison.png"),width=Inches(6.2)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
t=doc.add_table(rows=1,cols=7); t.style="Table Grid"
for c,x in zip(t.rows[0].cells,["模型","ROC-AUC (95% CI)","AUPRC (95% CI)","敏感度","特异度","F1","Brier"]): c.text=x
for model in names:
    row=[names[model],ci(m,"model",model,"roc_auc"),ci(m,"model",model,"auprc"),f"{val(m,'model',model,'sensitivity').estimate:.3f}",f"{val(m,'model',model,'specificity').estimate:.3f}",f"{val(m,'model',model,'f1').estimate:.3f}",f"{val(m,'model',model,'brier').estimate:.3f}"]
    for c,x in zip(t.add_row().cells,row): c.text=str(x)
doc.add_heading("4. 使用限制",1); doc.add_paragraph("这些数值只能用于预期结果展示、作图测试和论文版式预演。正式论文结果必须替换为真实队列在锁定分析流程下得到的OOF预测、置信区间和模型比较结果。")
doc.save(OUT); print(OUT)
