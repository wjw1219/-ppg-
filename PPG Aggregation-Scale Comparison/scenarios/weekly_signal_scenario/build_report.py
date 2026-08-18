import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
RES, FIG, REPORT = ROOT / "results", ROOT / "figures", ROOT / "report"
REPORT.mkdir(exist_ok=True)
summary = pd.read_csv(RES / "summary_metrics.csv")
paired = pd.read_csv(RES / "paired_differences.csv")
cost = pd.read_csv(RES / "computational_cost.csv")
meta = json.loads((ROOT / "data/scenario_metadata.json").read_text(encoding="utf-8"))
names = {"3day": "3天", "7day": "7天（预设主尺度）", "14day": "14天"}


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width)); tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths, font_size=8.3):
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW"); tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths))); tbl_w.set(qn("w:type"), "dxa")
    ind = OxmlElement("w:tblInd"); ind.set(qn("w:w"), "120"); ind.set(qn("w:type"), "dxa"); tbl_pr.append(ind)
    grid = table._tbl.tblGrid
    for node in list(grid): grid.remove(node)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row_index, row in enumerate(table.rows):
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width); cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_index == 0:
                shade = OxmlElement("w:shd"); shade.set(qn("w:fill"), "F2F4F7"); cell._tc.get_or_add_tcPr().append(shade)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0); paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"; run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(font_size); run.bold = row_index == 0


def metric(scale, name):
    return summary[(summary.scale == scale) & (summary.metric == name)].iloc[0]


def ci(scale, name):
    row = metric(scale, name)
    return f"{row.estimate:.3f} ({row.ci_lower:.3f}-{row.ci_upper:.3f})"


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5); section.page_height = Inches(11)
section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(0.492)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
for style_name, size, color, before, after in [
    ("Heading 1", 16, "2E74B5", 16, 8), ("Heading 2", 13, "2E74B5", 12, 6), ("Heading 3", 12, "1F4D78", 8, 4)
]:
    style = doc.styles[style_name]; style.font.name = "Calibri"
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size); style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]; header.text = "实验二 | 周尺度信号模拟情景"
for run in header.runs: run.font.size = Pt(9); run.font.color.rgb = RGBColor(100, 100, 100)
footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer.add_run("第 "); field = OxmlElement("w:fldSimple"); field.set(qn("w:instr"), "PAGE"); footer._p.append(field); footer.add_run(" 页")

title = doc.add_paragraph(); title.paragraph_format.space_before = Pt(16); title.paragraph_format.space_after = Pt(4)
run = title.add_run("实验二：PPG观测窗口聚合比较"); run.bold = True; run.font.size = Pt(23); run.font.color.rgb = RGBColor(31, 78, 121)
subtitle = doc.add_paragraph(); subtitle.paragraph_format.space_after = Pt(12)
run = subtitle.add_run("周尺度信号模拟情景 | 387例模拟患者 | 完整182天观测"); run.font.size = Pt(12); run.font.color.rgb = RGBColor(90, 90, 90)

warning = doc.add_table(rows=1, cols=1); warning.style = "Table Grid"
warning.cell(0, 0).text = "重要说明：本报告来自为检验方法而构造的模拟情景。7天信号幅度经过一次失败运行后的校准，结果不可作为真实临床证据或论文实证结果。"
style_table(warning, [9360], 9.5)

doc.add_heading("1. 实验设计", level=1)
doc.add_paragraph("本情景保持387例患者、房颤负荷增加标签、临床变量、患者级五折划分及模型结构不变，仅在模拟PPG特征中加入尺度相关信号。3天、7天和14天模型分别使用完整182天数据，并采用相同的5折交叉验证、3个训练随机种子和评价流程。")
doc.add_paragraph(f"预设信号幅度为3天 {meta['signal_amplitudes']['3day']:.2f}、7天 {meta['signal_amplitudes']['7day']:.2f}、14天 {meta['signal_amplitudes']['14day']:.2f}；对应不规则指数的标准化组间效应为 {meta['standardized_signal_effect']['3day']:.2f}、{meta['standardized_signal_effect']['7day']:.2f} 和 {meta['standardized_signal_effect']['14day']:.2f}。Holter标签及房颤负荷数值未被修改。")

doc.add_heading("2. 折外预测性能", level=1)
table = doc.add_table(rows=1, cols=8); table.style = "Table Grid"
headers = ["尺度", "AUROC (95% CI)", "AUPRC (95% CI)", "准确率 (95% CI)", "敏感度 (95% CI)", "特异度 (95% CI)", "F1 (95% CI)", "Brier (95% CI)"]
for cell, text in zip(table.rows[0].cells, headers): cell.text = text
for scale in ("3day", "7day", "14day"):
    values = [names[scale], ci(scale, "roc_auc"), ci(scale, "auprc"), ci(scale, "accuracy"), ci(scale, "sensitivity"), ci(scale, "specificity"), ci(scale, "f1"), ci(scale, "brier")]
    for cell, text in zip(table.add_row().cells, values): cell.text = text
style_table(table, [1100, 1260, 1260, 1180, 1120, 1120, 1160, 1160], 7.1)

doc.add_picture(str(FIG / "multiscale_roc.png"), width=Inches(5.8)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
caption = doc.add_paragraph("图1. 三种PPG聚合尺度的折外ROC曲线"); caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_picture(str(FIG / "multiscale_performance.png"), width=Inches(5.8)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
caption = doc.add_paragraph("图2. AUROC和AUPRC及患者级Bootstrap 95%置信区间"); caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading("3. 配对Bootstrap比较", level=1)
table = doc.add_table(rows=1, cols=3); table.style = "Table Grid"
for cell, text in zip(table.rows[0].cells, ["比较", "AUROC差值 (95% CI)", "AUPRC差值 (95% CI)"]): cell.text = text
for comparison, label in [("7day - 3day", "7天减3天"), ("7day - 14day", "7天减14天")]:
    values = paired[paired.comparison == comparison].set_index("metric")
    row = table.add_row().cells; row[0].text = label
    for cell, name in zip(row[1:], ["roc_auc_difference", "auprc_difference"]):
        value = values.loc[name]; cell.text = f"{value.estimate:+.3f} ({value.ci_lower:+.3f} 至 {value.ci_upper:+.3f})"
style_table(table, [2100, 3630, 3630], 9)

doc.add_heading("4. 计算成本", level=1)
table = doc.add_table(rows=1, cols=6); table.style = "Table Grid"
headers = ["尺度", "节点数", "平均epoch", "单次折-种子训练时间(s)", "每epoch时间(s)", "单例推理(ms)"]
for cell, text in zip(table.rows[0].cells, headers): cell.text = text
for scale in ("3day", "7day", "14day"):
    row_data = cost[cost.scale == scale].iloc[0]
    values = [names[scale], f"{int(row_data.nodes)}", f"{row_data.epochs_mean:.1f}", f"{row_data.training_seconds_mean:.2f}", f"{row_data.seconds_per_epoch_mean:.3f}", f"{row_data.inference_ms_per_patient_mean:.3f}"]
    for cell, text in zip(table.add_row().cells, values): cell.text = text
style_table(table, [1500, 1100, 1200, 2200, 1700, 1660], 8.2)
doc.add_picture(str(FIG / "performance_cost_tradeoff.png"), width=Inches(5.8)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
caption = doc.add_paragraph("图3. 折外AUROC与实际训练时间的关系；点面积表示时间节点数"); caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading("5. 结果解释", level=1)
doc.add_paragraph(f"本次校准情景中，7天模型的AUROC为{metric('7day','roc_auc').estimate:.3f}，低于3天的{metric('3day','roc_auc').estimate:.3f}和14天的{metric('14day','roc_auc').estimate:.3f}。因此，即使人工增强了7天尺度的原始信号，当前深度模型仍未在折外测试中形成7天优势。")
doc.add_paragraph("该结果表明，原始特征层的组间分离度不能保证时间图模型获得更好的泛化性能。继续根据测试结果修改数据或预测概率将构成结果导向调整。正式研究应保留7天预设主尺度，同时如实报告三个尺度及其置信区间，并在真实数据中验证。")

doc.add_heading("6. 可复现文件", level=1)
doc.add_paragraph("情景数据、生成元数据、折外预测、完整指标、配对差值、计算成本和图表均保存在 weekly_signal_scenario 目录。第一次校准失败运行保存在 calibration_run_1 子目录。")

output = REPORT / "实验二_周尺度信号模拟情景_结果报告.docx"
doc.save(output)
print(output)
