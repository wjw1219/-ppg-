from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from sklearn.metrics import accuracy_score, brier_score_loss, confusion_matrix, f1_score


ROOT = Path(__file__).resolve().parent
PREDICTIONS = ROOT / "results/oof_predictions.csv"
SUMMARY = ROOT / "results/summary_metrics.csv"
OUTPUT = ROOT / "report/实验二_各组混淆矩阵与分类指标核对.docx"
SCALES = [("3day", "3天"), ("7day", "7天（预设主尺度）"), ("14day", "14天")]


def set_cell_width(cell, width):
    prop = cell._tc.get_or_add_tcPr(); node = prop.find(qn("w:tcW"))
    if node is None: node = OxmlElement("w:tcW"); prop.append(node)
    node.set(qn("w:w"), str(width)); node.set(qn("w:type"), "dxa")


def style_table(table, widths, size=8.5, left_first=False):
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    prop = table._tbl.tblPr; node = prop.find(qn("w:tblW"))
    if node is None: node = OxmlElement("w:tblW"); prop.append(node)
    node.set(qn("w:w"), str(sum(widths))); node.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd"); indent.set(qn("w:w"), "120"); indent.set(qn("w:type"), "dxa"); prop.append(indent)
    grid = table._tbl.tblGrid
    for old in list(grid): grid.remove(old)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row_index, row in enumerate(table.rows):
        for column_index, (cell, width) in enumerate(zip(row.cells, widths)):
            set_cell_width(cell, width); cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_index == 0:
                fill = OxmlElement("w:shd"); fill.set(qn("w:fill"), "F2F4F7"); cell._tc.get_or_add_tcPr().append(fill)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if left_first and column_index == 0 and row_index else WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0); paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"; run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(size); run.bold = row_index == 0


predictions = pd.read_csv(PREDICTIONS)
summary = pd.read_csv(SUMMARY)
rows = []
for scale, display_name in SCALES:
    frame = predictions[predictions.scale.eq(scale)].sort_values("patient_id")
    y = frame.label.to_numpy(); probability = frame.probability.to_numpy()
    predicted = (probability >= frame.threshold.to_numpy()).astype(int)
    tn, fp, fn, tp = confusion_matrix(y, predicted, labels=[0, 1]).ravel()
    sensitivity = tp / (tp + fn); specificity = tn / (tn + fp)
    precision = tp / (tp + fp); npv = tn / (tn + fn)
    f1 = 2 * tp / (2 * tp + fp + fn); accuracy = (tp + tn) / len(y)
    rows.append({"scale": scale, "name": display_name, "n": len(y), "positive": int(y.sum()), "negative": int((y == 0).sum()),
                 "tp": int(tp), "fn": int(fn), "tn": int(tn), "fp": int(fp), "sensitivity": sensitivity,
                 "specificity": specificity, "precision": precision, "npv": npv, "f1": f1, "accuracy": accuracy,
                 "brier": brier_score_loss(y, probability), "sklearn_f1": f1_score(y, predicted),
                 "sklearn_accuracy": accuracy_score(y, predicted)})
audit = pd.DataFrame(rows)

for row in rows:
    for metric in ("sensitivity", "specificity", "f1", "accuracy", "brier"):
        reported = summary[(summary.scale.eq(row["scale"])) & (summary.metric.eq(metric))].estimate.iloc[0]
        if abs(reported - row[metric]) > 1e-12:
            raise ValueError(f"Metric mismatch: {row['scale']} {metric}: {row[metric]} != {reported}")
    if abs(row["f1"] - row["sklearn_f1"]) > 1e-12 or abs(row["accuracy"] - row["sklearn_accuracy"]) > 1e-12:
        raise ValueError(f"Independent sklearn verification failed for {row['scale']}")

audit.to_csv(ROOT / "results/confusion_matrix_and_metric_audit.csv", index=False, encoding="utf-8-sig")

doc = Document(); section = doc.sections[0]
section.page_width = Inches(8.5); section.page_height = Inches(11)
section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(0.492)
normal = doc.styles["Normal"]; normal.font.name = "Calibri"
normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in [("Heading 1", 16, "2E74B5", 16, 8), ("Heading 2", 13, "2E74B5", 12, 6)]:
    style = doc.styles[name]; style.font.name = "Calibri"; style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size); style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after); style.paragraph_format.keep_with_next = True
header = section.header.paragraphs[0]; header.text = "实验二 | 混淆矩阵与分类指标核对"
for run in header.runs: run.font.size = Pt(9); run.font.color.rgb = RGBColor(100, 100, 100)
footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer.add_run("第 "); page = OxmlElement("w:fldSimple"); page.set(qn("w:instr"), "PAGE"); footer._p.append(page); footer.add_run(" 页")

title = doc.add_paragraph(); title.paragraph_format.space_before = Pt(16); title.paragraph_format.space_after = Pt(4)
run = title.add_run("实验二：各组混淆矩阵与分类指标核对"); run.bold = True; run.font.size = Pt(22); run.font.color.rgb = RGBColor(31, 78, 121)
subtitle = doc.add_paragraph(); subtitle.paragraph_format.space_after = Pt(12)
run = subtitle.add_run("基于当前模拟OOF预测 | 387例患者 | 分类阈值0.5"); run.font.size = Pt(12); run.font.color.rgb = RGBColor(90, 90, 90)
notice = doc.add_table(rows=1, cols=1); notice.style = "Table Grid"
notice.cell(0, 0).text = "说明：当前预测概率为模拟结果，并非真实模型训练输出。本文件用于核对混淆矩阵和评价公式。"
style_table(notice, [9360], 9.5)

doc.add_heading("1. 各组混淆矩阵", level=1)
table = doc.add_table(rows=1, cols=8); table.style = "Table Grid"
for cell, text in zip(table.rows[0].cells, ["尺度", "总例数", "实际阳性", "实际阴性", "TP", "FN", "TN", "FP"]): cell.text = text
for row in rows:
    values = [row["name"], row["n"], row["positive"], row["negative"], row["tp"], row["fn"], row["tn"], row["fp"]]
    for cell, value in zip(table.add_row().cells, values): cell.text = str(value)
style_table(table, [1800, 1080, 1080, 1080, 1080, 1080, 1080, 1080], 8.5)

doc.add_heading("2. 指标定义", level=1)
for text in ["敏感度 = TP / (TP + FN)", "特异度 = TN / (TN + FP)", "精确率 = TP / (TP + FP)",
             "阴性预测值 = TN / (TN + FN)", "F1 = 2TP / (2TP + FP + FN)", "准确率 = (TP + TN) / (TP + TN + FP + FN)",
             "Brier = (1/N) × Σ(预测概率 - 真实标签)²"]:
    doc.add_paragraph(text)

doc.add_heading("3. 公式代入与结果", level=1)
table = doc.add_table(rows=1, cols=7); table.style = "Table Grid"
for cell, text in zip(table.rows[0].cells, ["尺度", "敏感度", "特异度", "精确率", "阴性预测值", "F1", "准确率"]): cell.text = text
for row in rows:
    values = [row["name"], f"{row['tp']}/({row['tp']}+{row['fn']})={row['sensitivity']:.6f}",
              f"{row['tn']}/({row['tn']}+{row['fp']})={row['specificity']:.6f}",
              f"{row['tp']}/({row['tp']}+{row['fp']})={row['precision']:.6f}",
              f"{row['tn']}/({row['tn']}+{row['fn']})={row['npv']:.6f}",
              f"2×{row['tp']}/(2×{row['tp']}+{row['fp']}+{row['fn']})={row['f1']:.6f}",
              f"({row['tp']}+{row['tn']})/{row['n']}={row['accuracy']:.6f}"]
    for cell, value in zip(table.add_row().cells, values): cell.text = value
style_table(table, [1450, 1400, 1400, 1280, 1280, 1360, 1190], 7.1)

doc.add_heading("4. 与结果表的独立核对", level=1)
table = doc.add_table(rows=1, cols=7); table.style = "Table Grid"
for cell, text in zip(table.rows[0].cells, ["尺度", "敏感度", "特异度", "F1", "准确率", "Brier", "核对结论"]): cell.text = text
for row in rows:
    values = [row["name"], f"{row['sensitivity']:.3f}", f"{row['specificity']:.3f}", f"{row['f1']:.3f}",
              f"{row['accuracy']:.3f}", f"{row['brier']:.3f}", "与summary_metrics.csv一致"]
    for cell, value in zip(table.add_row().cells, values): cell.text = value
style_table(table, [1700, 1150, 1150, 1050, 1050, 1050, 2210], 8.2)

doc.add_paragraph("核对结果：三组的敏感度、特异度、F1、准确率和Brier均由当前OOF预测重新计算，与summary_metrics.csv中的点估计在1×10^-12容差内完全一致。F1和准确率同时使用scikit-learn独立复算，结果一致。")
OUTPUT.parent.mkdir(exist_ok=True); doc.save(OUTPUT); print(OUTPUT)
