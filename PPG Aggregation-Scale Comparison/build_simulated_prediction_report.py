import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
RES, FIG, REPORT = ROOT / "results", ROOT / "figures", ROOT / "report"
summary = pd.read_csv(RES / "summary_metrics.csv")
paired = pd.read_csv(RES / "paired_differences.csv")
metadata = json.loads((RES / "simulated_prediction_metadata.json").read_text(encoding="utf-8"))
names = {"3day": "3天", "7day": "7天（预设主尺度）", "14day": "14天"}


def cell_width(cell, width):
    prop = cell._tc.get_or_add_tcPr(); node = prop.find(qn("w:tcW"))
    if node is None: node = OxmlElement("w:tcW"); prop.append(node)
    node.set(qn("w:w"), str(width)); node.set(qn("w:type"), "dxa")


def style_table(table, widths, size=8):
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    prop = table._tbl.tblPr; width_node = prop.find(qn("w:tblW"))
    if width_node is None: width_node = OxmlElement("w:tblW"); prop.append(width_node)
    width_node.set(qn("w:w"), str(sum(widths))); width_node.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd"); indent.set(qn("w:w"), "120"); indent.set(qn("w:type"), "dxa"); prop.append(indent)
    grid = table._tbl.tblGrid
    for node in list(grid): grid.remove(node)
    for width in widths:
        node = OxmlElement("w:gridCol"); node.set(qn("w:w"), str(width)); grid.append(node)
    for index, row in enumerate(table.rows):
        for cell, width in zip(row.cells, widths):
            cell_width(cell, width); cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if index == 0:
                fill = OxmlElement("w:shd"); fill.set(qn("w:fill"), "F2F4F7"); cell._tc.get_or_add_tcPr().append(fill)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_before = Pt(0); paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"; run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(size); run.bold = index == 0


def value(scale, metric):
    return summary[(summary.scale == scale) & (summary.metric == metric)].iloc[0]


def ci(scale, metric):
    row = value(scale, metric)
    return f"{row.estimate:.3f} ({row.ci_lower:.3f}-{row.ci_upper:.3f})"


doc = Document(); section = doc.sections[0]
section.page_width = Inches(8.5); section.page_height = Inches(11)
section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(0.492)
normal = doc.styles["Normal"]; normal.font.name = "Calibri"
normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in [("Heading 1", 16, "2E74B5", 16, 8), ("Heading 2", 13, "2E74B5", 12, 6)]:
    style = doc.styles[name]; style.font.name = "Calibri"; style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size); style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after); style.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]; header.text = "实验二 | 模拟折外预测结果"
for run in header.runs: run.font.size = Pt(9); run.font.color.rgb = RGBColor(100, 100, 100)
footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer.add_run("第 "); page = OxmlElement("w:fldSimple"); page.set(qn("w:instr"), "PAGE"); footer._p.append(page); footer.add_run(" 页")

title = doc.add_paragraph(); title.paragraph_format.space_before = Pt(16); title.paragraph_format.space_after = Pt(4)
run = title.add_run("实验二：观测窗口聚合比较"); run.bold = True; run.font.size = Pt(23); run.font.color.rgb = RGBColor(31, 78, 121)
subtitle = doc.add_paragraph(); subtitle.paragraph_format.space_after = Pt(12)
run = subtitle.add_run("替换后的模拟OOF预测 | 387例患者 | 3天、7天和14天聚合"); run.font.size = Pt(12); run.font.color.rgb = RGBColor(90, 90, 90)

warning = doc.add_table(rows=1, cols=1); warning.style = "Table Grid"
warning.cell(0, 0).text = "重要说明：本报告中的预测概率为模拟生成，并非模型重新训练所得，仅用于分析流程、表格和图形展示，不得作为真实临床研究结果。"
style_table(warning, [9360], 9.5)

doc.add_heading("1. 模拟预测设置", level=1)
doc.add_paragraph(f"患者ID、房颤负荷增加标签和三个比较尺度保持不变。预测概率使用固定随机种子 {metadata['seed']} 生成，分类阈值固定为0.5。7天条件预设188例真阳性、12例假阴性、176例真阴性和11例假阳性；概率分布经固定参数校准，使AUROC、AUPRC和Brier接近预设目标。")
doc.add_paragraph("全部性能指标均由替换后的1,161条折外预测重新计算。95%置信区间使用2,000次患者级分层Bootstrap；7天与其他尺度的AUROC和AUPRC差值使用相同重采样索引进行配对计算。")

doc.add_heading("2. 重新计算的性能指标", level=1)
table = doc.add_table(rows=1, cols=8); table.style = "Table Grid"
headers = ["尺度", "AUROC (95% CI)", "AUPRC (95% CI)", "准确率 (95% CI)", "敏感度 (95% CI)", "特异度 (95% CI)", "F1 (95% CI)", "Brier (95% CI)"]
for cell, text in zip(table.rows[0].cells, headers): cell.text = text
for scale in ("3day", "7day", "14day"):
    values = [names[scale], ci(scale, "roc_auc"), ci(scale, "auprc"), ci(scale, "accuracy"), ci(scale, "sensitivity"), ci(scale, "specificity"), ci(scale, "f1"), ci(scale, "brier")]
    for cell, text in zip(table.add_row().cells, values): cell.text = text
style_table(table, [1100, 1260, 1260, 1180, 1120, 1120, 1160, 1160], 7.1)

doc.add_picture(str(FIG / "multiscale_roc.png"), width=Inches(5.8)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
caption = doc.add_paragraph("图1. 三种聚合尺度的模拟折外ROC曲线"); caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_picture(str(FIG / "multiscale_performance.png"), width=Inches(5.8)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
caption = doc.add_paragraph("图2. AUROC和AUPRC及Bootstrap 95%置信区间"); caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading("3. 配对比较", level=1)
table = doc.add_table(rows=1, cols=3); table.style = "Table Grid"
for cell, text in zip(table.rows[0].cells, ["比较", "AUROC差值 (95% CI)", "AUPRC差值 (95% CI)"]): cell.text = text
for comparison, label in [("7day - 3day", "7天减3天"), ("7day - 14day", "7天减14天")]:
    frame = paired[paired.comparison == comparison].set_index("metric"); row = table.add_row().cells; row[0].text = label
    for cell, metric in zip(row[1:], ["roc_auc_difference", "auprc_difference"]):
        item = frame.loc[metric]; cell.text = f"{item.estimate:+.3f} ({item.ci_lower:+.3f} 至 {item.ci_upper:+.3f})"
style_table(table, [2100, 3630, 3630], 9)

doc.add_heading("4. 结果概述", level=1)
doc.add_paragraph("7天尺度在AUROC、AUPRC、准确率、敏感度、特异度和F1上均为最高，同时Brier score最低，表示其模拟预测具有最好的区分度、分类性能和概率准确性。7天相对3天的AUROC与AUPRC配对差值置信区间未跨越0；相对14天的差值点估计为正，但置信区间跨越0。")
doc.add_paragraph("上述描述仅适用于本次模拟预测文件。正式论文必须使用真实模型训练产生的折外预测重新计算，并删除模拟结果声明后方可作为实证结果报告。")

REPORT.mkdir(exist_ok=True)
output = REPORT / "实验二_模拟OOF预测替换后_结果报告.docx"
doc.save(output); print(output)
