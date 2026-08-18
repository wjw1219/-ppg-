import json,unittest
from pathlib import Path
import pandas as pd
from docx import Document
ROOT=Path(__file__).resolve().parents[1]

@unittest.skipUnless((ROOT/"results/oof_predictions.csv").exists(), "Run the experiment before output-structure checks")
class GeneratedOutputTests(unittest.TestCase):
    pass

class OutputTests(GeneratedOutputTests):
    def test_prediction_outputs(self):
        p=pd.read_csv(ROOT/"results/oof_predictions.csv"); self.assertEqual(len(p),1161); self.assertEqual(p.scale.nunique(),3); self.assertTrue(p.groupby(["patient_id","scale"]).size().eq(1).all()); self.assertTrue(p.probability.between(0,1).all())
        self.assertTrue(all(json.loads((ROOT/"results/quality_checks.json").read_text()).values()))
    def test_report_structure(self):
        d=Document(ROOT/"report/实验二_观测窗口聚合比较_实验说明.docx"); text="\n".join(x.text for x in d.paragraphs)
        self.assertEqual(len(d.inline_shapes),2); self.assertEqual(len(d.tables),3); self.assertEqual(sum(x.style.name=="Heading 1" for x in d.paragraphs),11); self.assertIn("7天为预设主分析尺度",text); self.assertNotIn("TODO",text)
        self.assertAlmostEqual(d.sections[0].left_margin.inches,1,places=2); self.assertAlmostEqual(d.sections[0].right_margin.inches,1,places=2)

if __name__=="__main__": unittest.main()
