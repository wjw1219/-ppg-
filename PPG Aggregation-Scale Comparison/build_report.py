from pathlib import Path
import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT=Path(__file__).resolve().parent; RES=ROOT/"results"; FIG=ROOT/"figures"; OUT=ROOT/"report/实验二_观测窗口聚合比较_实验说明.docx"
summary=pd.read_csv(RES/"summary_metrics.csv"); paired=pd.read_csv(RES/"paired_differences.csv")
NAMES={"3day":"3天","7day":"7天（预设主尺度）","14day":"14天"}

def shade(cell,fill):
    x=OxmlElement("w:shd"); x.set(qn("w:fill"),fill); cell._tc.get_or_add_tcPr().append(x)
def set_width(cell,w):
    pr=cell._tc.get_or_add_tcPr(); x=pr.find(qn("w:tcW"))
    if x is None: x=OxmlElement("w:tcW"); pr.append(x)
    x.set(qn("w:w"),str(w)); x.set(qn("w:type"),"dxa")
def style_table(t,widths,size=8.5):
    t.autofit=False; t.alignment=WD_TABLE_ALIGNMENT.CENTER; pr=t._tbl.tblPr; tw=pr.find(qn("w:tblW"))
    if tw is None: tw=OxmlElement("w:tblW"); pr.append(tw)
    tw.set(qn("w:w"),str(sum(widths))); tw.set(qn("w:type"),"dxa"); ind=OxmlElement("w:tblInd"); ind.set(qn("w:w"),"120"); ind.set(qn("w:type"),"dxa"); pr.append(ind)
    grid=t._tbl.tblGrid
    for x in list(grid): grid.remove(x)
    for w in widths: x=OxmlElement("w:gridCol"); x.set(qn("w:w"),str(w)); grid.append(x)
    for i,row in enumerate(t.rows):
        for c,w in zip(row.cells,widths):
            set_width(c,w); c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
            if i==0: shade(c,"F2F4F7")
            for p in c.paragraphs:
                p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
                for r in p.runs: r.font.name="Calibri"; r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"),"Microsoft YaHei"); r.font.size=Pt(size); r.bold=i==0
def row(scale,metric): return summary[(summary.scale==scale)&(summary.metric==metric)].iloc[0]
def ci(scale,metric):
    r=row(scale,metric); return f"{r.estimate:.3f} ({r.ci_lower:.3f}-{r.ci_upper:.3f})"

doc=Document(); sec=doc.sections[0]; sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1); sec.header_distance=sec.footer_distance=Inches(.492)
n=doc.styles["Normal"]; n.font.name="Calibri"; n._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"),"Microsoft YaHei"); n.font.size=Pt(11); n.paragraph_format.space_after=Pt(6); n.paragraph_format.line_spacing=1.10
for name,size,color,before,after in [("Heading 1",16,"2E74B5",16,8),("Heading 2",13,"2E74B5",12,6),("Heading 3",12,"1F4D78",8,4)]:
    s=doc.styles[name]; s.font.name="Calibri"; s._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"),"Microsoft YaHei"); s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after); s.paragraph_format.keep_with_next=True
hdr=sec.header.paragraphs[0]; hdr.text="实验二 | 观测窗口聚合比较"
for r in hdr.runs: r.font.size=Pt(9); r.font.color.rgb=RGBColor(100,100,100)
f=sec.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.RIGHT; f.add_run("第 "); fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),"PAGE"); f._p.append(fld); f.add_run(" 页")
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(16); p.paragraph_format.space_after=Pt(4); r=p.add_run("实验二：观测窗口聚合比较"); r.bold=True; r.font.size=Pt(23); r.font.color.rgb=RGBColor(31,78,121); r.font.name="Calibri"; r._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"),"Microsoft YaHei")
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(14); r=p.add_run("3天、7天和14天PPG聚合 | 完整26周观察 | 固定融合模型"); r.font.size=Pt(12); r.font.color.rgb=RGBColor(90,90,90)
for a,b in [("分析对象","387例患者"),("预设主尺度","7天"),("比较尺度","3天、14天"),("主要指标","ROC-AUC"),("运行日期","2026年7月23日")]: p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2); x=p.add_run(a+"："); x.bold=True; p.add_run(b)

doc.add_heading("1. 实验目的",1); doc.add_paragraph("本实验保持个人先验信息、表示层融合模型和完整182天观察范围不变，比较3天、7天和14天PPG聚合对纵向表征与房颤负荷增加预测性能的影响。7天为预设主分析尺度，比较结果不用于根据测试集表现事后改变主尺度。")
doc.add_heading("2. 数据与结局",1); doc.add_paragraph("本实验使用387例模拟患者数据验证多尺度比较流程。出院和6个月72小时Holter房颤负荷仅用于确定负荷增加标签，不进入临床分支、PPG分支或融合分类器。200例（51.7%）患者负荷增加，187例（48.3%）未增加。")
doc.add_heading("3. 聚合尺度",1)
t=doc.add_table(rows=1,cols=5); t.style="Table Grid"
for c,h in zip(t.rows[0].cells,["尺度","时间节点","完整窗口","末窗口","研究定位"]): c.text=h
for values in [("3天","61","60个3天窗口","2天","比较尺度"),("7天","26","26个7天窗口","7天","预设主尺度"),("14天","13","13个14天窗口","14天","比较尺度")]:
    for c,v in zip(t.add_row().cells,values): c.text=v
style_table(t,[1300,1500,2100,1600,2860],9)
doc.add_paragraph("三个尺度均覆盖第1至182天，直接使用真实聚合序列，不进行插值、下采样或统一节点数。每个时间图仅设置相邻窗口双向边及自环。")
doc.add_heading("4. 模型与公平比较",1); doc.add_paragraph("三种尺度使用相同的完整临床先验编码器。PPG分支分别处理61、26和13个时间节点，通过相邻窗口图注意力及掩码注意力池化得到患者级PPG表示。临床表示与PPG表示直接拼接后输入相同的两层MLP分类器。除序列长度及位置编码长度外，模型结构与超参数保持一致。")
doc.add_paragraph("所有尺度共享同一患者级分层五折、20%内部验证比例、3个训练随机种子、早停规则、隐藏维度、注意力头数、学习率、权重衰减和批大小。每个尺度的填补与标准化参数仅在相应训练数据中计算。")
doc.add_heading("5. 评价与统计",1); doc.add_paragraph("ROC-AUC为主要指标，同时报告AUPRC、敏感度、特异度、F1、准确率和Brier score。分类阈值由内部验证集Youden指数确定。各尺度置信区间及7天相对3天、14天的性能差值使用相同的2,000组分层患者级Bootstrap索引计算。")

doc.add_heading("6. 三尺度预测性能",1)
t=doc.add_table(rows=1,cols=7); t.style="Table Grid"
for c,h in zip(t.rows[0].cells,["尺度","ROC-AUC (95% CI)","AUPRC (95% CI)","敏感度","特异度","F1","Brier"]): c.text=h
for scale in ("3day","7day","14day"):
    vals=[NAMES[scale],ci(scale,"roc_auc"),ci(scale,"auprc"),f"{row(scale,'sensitivity').estimate:.3f}",f"{row(scale,'specificity').estimate:.3f}",f"{row(scale,'f1').estimate:.3f}",f"{row(scale,'brier').estimate:.3f}"]
    for c,v in zip(t.add_row().cells,vals): c.text=v
style_table(t,[1500,1650,1650,1100,1100,1000,1360],8)
doc.add_paragraph("14天聚合的ROC-AUC点估计最高（0.619），其次为7天（0.572）和3天（0.535）。AUPRC呈相同排序。该排序仅描述当前折外预测结果，不能据此事后将14天改为主分析尺度。")
doc.add_picture(str(FIG/"multiscale_roc.png"),width=Inches(6.1)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
p=doc.add_paragraph("图1. 3天、7天和14天融合模型的折外ROC曲线"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].italic=True

doc.add_heading("7. 相对7天主尺度的配对差值",1)
t=doc.add_table(rows=1,cols=3); t.style="Table Grid"
for c,h in zip(t.rows[0].cells,["比较","ROC-AUC差值 (95% CI)","AUPRC差值 (95% CI)"]): c.text=h
for comp,label in [("7day - 3day","7天减3天"),("7day - 14day","7天减14天")]:
    x=paired[paired.comparison==comp].set_index("metric"); cells=t.add_row().cells; cells[0].text=label
    for c,m in zip(cells[1:],["roc_auc_difference","auprc_difference"]): r=x.loc[m]; c.text=f"{r.estimate:+.3f} ({r.ci_lower:+.3f}至{r.ci_upper:+.3f})"
style_table(t,[2500,3430,3430],9)
doc.add_paragraph("7天相对3天的ROC-AUC差值为+0.037（95% CI -0.021至0.098）；7天相对14天的差值为-0.046（95% CI -0.101至0.007）。两项区间均跨0，没有证据表明7天与比较尺度之间存在稳定差异。")
doc.add_picture(str(FIG/"multiscale_performance.png"),width=Inches(6.1)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
p=doc.add_paragraph("图2. 三种聚合尺度的ROC-AUC和AUPRC及95%置信区间"); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.runs[0].italic=True

doc.add_heading("8. 结果解释",1); doc.add_paragraph("较短窗口保留更细粒度变化，但增加时间节点数和局部噪声；较长窗口提高单个节点的稳定性，但平滑短时波动。本实验中14天尺度点估计较高，可能反映更强的降噪效应。然而，配对置信区间仍跨0，不能认定14天稳定优于7天，也不能将预测差异解释为生理时间尺度的因果效应。")
doc.add_heading("9. 预设主尺度的处理",1); doc.add_paragraph("7天在实验前已被指定为主分析尺度。本实验属于多尺度聚合比较，3天和14天结果用于评价尺度敏感性。正式报告应同时给出三种结果，但后续主要分析仍使用7天，不依据本次测试折表现重新选择聚合尺度。")
doc.add_heading("10. 实验限制",1)
for x in ["样本量为387例，三个深度模型的置信区间较宽。","不同尺度的节点数不同，模型参数量相近但计算路径长度不同。","最后一个3天窗口只覆盖2天，其记录数和稳定性可能与完整3天窗口不同。","本实验固定完整182天输入，不能说明不同观察时长下最适合的聚合尺度。"]: doc.add_paragraph(x,style="List Bullet")
doc.add_heading("11. 可复现输出",1); doc.add_paragraph("实验目录保存逐患者折外预测、逐折训练记录、总体指标、配对差值、质量检查和两张结果图。所有表格均可从results/oof_predictions.csv重新计算。")
doc.save(OUT); print(OUT)
